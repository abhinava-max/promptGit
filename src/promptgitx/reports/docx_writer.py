"""Create polished DOCX review reports."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


COLOR_NAVY = "1E3A5F"
COLOR_BLUE = "2563EB"
COLOR_MUTED = "64748B"
COLOR_FILE = "475569"
COLOR_GREEN = "15803D"
COLOR_RED = "B91C1C"
COLOR_YELLOW = "B45309"
COLOR_DARK = "111827"
COLOR_HEADER_BG = "E8EEF6"
COLOR_CARD_BORDER = "D7DEE8"
COLOR_CARD_BG = "FBFCFE"
COLOR_SEVERITY_MEDIUM_BG = "FEF3C7"
COLOR_SEVERITY_HIGH_BG = "FEE2E2"
COLOR_SEVERITY_LOW_BG = "EAF7EF"


def findings_by_id(report: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {
        str(finding.get("id")): finding
        for finding in report.get("findings", [])
        if finding.get("id")
    }


def resolve_issue_items(report: dict[str, Any], items: list[Any]) -> list[dict[str, Any]]:
    lookup = findings_by_id(report)
    resolved = []

    for item in items:
        if isinstance(item, str):
            finding = lookup.get(item)

            if finding:
                resolved.append(finding)
        elif isinstance(item, dict):
            resolved.append(item)

    return resolved


def get_file_issues(report: dict[str, Any], file: dict[str, Any]) -> list[dict[str, Any]]:
    if file.get("issues"):
        return file.get("issues", [])

    return resolve_issue_items(report, file.get("issue_ids", []))


def group_items_by_file(items: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped: dict[str, list[dict[str, Any]]] = {}

    for item in items:
        file_path = str(item.get("file_path", "unknown"))
        grouped.setdefault(file_path, []).append(item)

    return grouped


def issue_message(item: dict[str, Any]) -> str:
    return str(item.get("short_message") or item.get("message") or item.get("suggestion") or "No details provided.")


def issue_ids_text(items: list[dict[str, Any]]) -> str:
    issue_ids = [str(item["id"]) for item in items if item.get("id")]

    if not issue_ids:
        return ""

    return ", ".join(issue_ids)


def format_line_reference(reference: Any) -> str:
    reference_text = str(reference or "").strip()

    if not reference_text:
        return ""

    lowered = reference_text.lower()

    if "changed file" in lowered:
        return "Changed file"

    line_matches = re.findall(r"(?:old\s+)?L(\d+)(?:-(?:old\s+)?L?(\d+))?", reference_text)

    if not line_matches:
        return reference_text

    formatted = []

    for start, end in line_matches:
        if end:
            formatted.append(f"Line {start}:{end}")
        else:
            formatted.append(f"Line {start}")

    return ", ".join(formatted)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = "E5E7EB"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_paragraph(paragraph, size: int = 10, color: str = COLOR_DARK, bold: bool = False):
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_colored_run(paragraph, text: str, color: str, bold: bool = False, size: int = 10):
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    return run


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_colored_run(paragraph, "Page ", COLOR_MUTED, size=8)

    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLOR_MUTED)


def severity_color(severity: str | None) -> str:
    severity_value = str(severity or "").upper()

    if severity_value in {"CRITICAL", "HIGH"}:
        return COLOR_RED
    if severity_value == "MEDIUM":
        return COLOR_YELLOW
    return COLOR_GREEN


def severity_background(severity: str | None) -> str:
    severity_value = str(severity or "").upper()

    if severity_value in {"CRITICAL", "HIGH"}:
        return COLOR_SEVERITY_HIGH_BG
    if severity_value == "MEDIUM":
        return COLOR_SEVERITY_MEDIUM_BG
    return COLOR_SEVERITY_LOW_BG


def recommendation_color(recommendation: str | None) -> str:
    recommendation_value = str(recommendation or "").upper()

    if recommendation_value == "REQUEST_CHANGES":
        return COLOR_RED
    if recommendation_value == "NEEDS_MANUAL_REVIEW":
        return COLOR_YELLOW
    return COLOR_GREEN


def title_case(value: str | None) -> str:
    if not value:
        return "Unknown"

    return str(value).replace("_", " ").title()


def add_heading(document: Document, text: str, level: int = 1):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for run in heading.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(COLOR_NAVY if level == 1 else COLOR_BLUE)

    return heading


def add_key_value_table(document: Document, rows: list[tuple[str, Any, str | None]]):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for label, value, color in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
        set_cell_shading(row.cells[0], COLOR_HEADER_BG)
        set_cell_border(row.cells[0])
        set_cell_border(row.cells[1])
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style_paragraph(row.cells[0].paragraphs[0], bold=True, color=COLOR_MUTED)
        style_paragraph(row.cells[1].paragraphs[0], bold=bool(color), color=color or COLOR_DARK)

    return table


def add_files_table(document: Document, files: list[dict[str, Any]]):
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["File", "Added", "Removed"]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, COLOR_HEADER_BG)
        set_cell_border(cell)
        style_paragraph(cell.paragraphs[0], bold=True, color=COLOR_NAVY)

    for file in files:
        row = table.add_row()
        row.cells[0].text = str(file.get("file_path", "unknown"))
        row.cells[1].text = f"+{file.get('added_lines', 0)}"
        row.cells[2].text = f"-{file.get('removed_lines', 0)}"

        set_cell_border(row.cells[0])
        set_cell_border(row.cells[1])
        set_cell_border(row.cells[2])
        style_paragraph(row.cells[0].paragraphs[0], color=COLOR_FILE, bold=True)
        style_paragraph(row.cells[1].paragraphs[0], color=COLOR_GREEN, bold=True)
        style_paragraph(row.cells[2].paragraphs[0], color=COLOR_RED, bold=True)


def add_executive_summary(document: Document, report: dict[str, Any]):
    summary = report.get("summary", {})
    recommendation = summary.get("final_recommendation", "APPROVE")
    risk = summary.get("risk_level", "low")
    issue_count = len(report.get("findings", [])) or sum(len(file.get("issues", [])) for file in report.get("files", []))

    add_heading(document, "Executive Summary", level=1)
    table = document.add_table(rows=0, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.add_row()
    cell = row.cells[0]
    set_cell_border(cell, COLOR_CARD_BORDER)
    set_cell_shading(cell, COLOR_CARD_BG)

    headline = cell.paragraphs[0]
    add_colored_run(headline, f"{title_case(risk)} risk", severity_color(risk), bold=True, size=13)
    add_colored_run(headline, " · ", COLOR_MUTED, size=13)
    add_colored_run(headline, str(recommendation), recommendation_color(recommendation), bold=True, size=13)

    body = cell.add_paragraph()
    add_colored_run(
        body,
        f"Reviewed {summary.get('files_reviewed', 0)} file(s), "
        f"{summary.get('lines_added', 0)} line(s) added, "
        f"{summary.get('lines_removed', 0)} line(s) removed, "
        f"and {issue_count} issue(s) flagged.",
        COLOR_DARK,
    )

    reason = cell.add_paragraph()
    add_colored_run(reason, str(report.get("overall_summary", "No summary available.")), COLOR_MUTED)


def add_issue_section(document: Document, title: str, items: list[dict[str, Any]]):
    if not items:
        return

    add_heading(document, title, level=2)

    for file_path, file_items in group_items_by_file(items).items():
        table = document.add_table(rows=0, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        row = table.add_row()
        cell = row.cells[0]
        set_cell_border(cell, COLOR_CARD_BORDER)
        set_cell_shading(cell, COLOR_CARD_BG)
        paragraph = cell.paragraphs[0]
        add_colored_run(paragraph, file_path, COLOR_FILE, bold=True)

        for item in file_items:
            issue = cell.add_paragraph()
            issue.paragraph_format.left_indent = Inches(0.25)
            add_colored_run(issue, "- ", COLOR_MUTED, bold=True)
            add_colored_run(issue, issue_message(item), COLOR_DARK)

            if item.get("line_reference"):
                reference = cell.add_paragraph()
                reference.paragraph_format.left_indent = Inches(0.45)
                add_colored_run(reference, "refs: ", COLOR_MUTED, bold=True)
                add_colored_run(reference, format_line_reference(item["line_reference"]), COLOR_DARK)

        ids = issue_ids_text(file_items)

        if ids:
            id_line = cell.add_paragraph()
            id_line.paragraph_format.left_indent = Inches(0.25)
            add_colored_run(id_line, "- ids: ", COLOR_MUTED, bold=True)
            add_colored_run(id_line, ids, COLOR_DARK)

        document.add_paragraph()


def add_file_wise_issues(document: Document, report: dict[str, Any]):
    add_heading(document, "File-wise Issues", level=2)
    files = report.get("files", [])
    found_issue = False

    for file in files:
        for issue in get_file_issues(report, file):
            found_issue = True
            table = document.add_table(rows=0, cols=1)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            row = table.add_row()
            cell = row.cells[0]
            set_cell_border(cell, COLOR_CARD_BORDER)
            paragraph = cell.paragraphs[0]
            severity = str(issue.get("severity", "low")).upper()
            if issue.get("id"):
                add_colored_run(paragraph, f"{issue['id']}  ", COLOR_MUTED, bold=True)
            add_colored_run(paragraph, f"{severity}  ", severity_color(severity), bold=True)
            add_colored_run(paragraph, title_case(issue.get("category")), COLOR_NAVY, bold=True)
            set_cell_shading(cell, severity_background(severity))

            details = [
                ("File", file.get("file_path", "unknown")),
                ("Problem", issue.get("message", "No details provided.")),
                ("Reference", format_line_reference(issue.get("line_reference", ""))),
                ("Suggestion", issue.get("suggestion", "")),
            ]

            for label, value in details:
                if not value:
                    continue

                detail = cell.add_paragraph()
                add_colored_run(detail, f"{label}: ", COLOR_MUTED, bold=True)
                color = COLOR_FILE if label == "File" else COLOR_DARK
                add_colored_run(detail, str(value), color, bold=(label == "File"))

            document.add_paragraph()

    if not found_issue:
        paragraph = document.add_paragraph()
        add_colored_run(paragraph, "None found.", COLOR_MUTED)


def add_clean_checks(document: Document, issues: dict[str, list[dict[str, Any]]]):
    checks = [
        ("Critical", issues.get("critical", [])),
        ("Breaking", issues.get("breaking_changes", [])),
        ("Security", issues.get("security", [])),
        ("Unprofessional Language", issues.get("vulgarity", [])),
        ("Performance", issues.get("performance", [])),
    ]
    clean = [name for name, items in checks if not items]

    add_heading(document, "Clean Checks", level=2)
    paragraph = document.add_paragraph()

    if clean:
        add_colored_run(paragraph, "No findings in: ", COLOR_MUTED, bold=True)
        add_colored_run(paragraph, ", ".join(clean), COLOR_FILE, bold=True)
    else:
        add_colored_run(paragraph, "Issues found in all major check groups.", COLOR_YELLOW, bold=True)


def create_docx_report(report: dict[str, Any], output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.start_type = WD_SECTION_START.NEW_PAGE
    add_page_number(section)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_colored_run(title, "GitPromptX Review Report", COLOR_NAVY, bold=True, size=24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_colored_run(subtitle, "AI-powered Git review summary", COLOR_MUTED, size=11)

    document.add_paragraph()

    target = report.get("target", {})
    target_value = target.get("value", "")

    if isinstance(target_value, list):
        target_value = ", ".join(str(item) for item in target_value)

    add_executive_summary(document, report)

    add_heading(document, "Review Target", level=1)
    add_key_value_table(
        document,
        [
            ("Mode", f"{title_case(target.get('mode'))} Review", None),
            ("Target", target_value, None),
        ],
    )

    summary = report.get("summary", {})
    recommendation = summary.get("final_recommendation", "APPROVE")
    risk = summary.get("risk_level", "low")
    add_heading(document, "Summary", level=1)
    add_key_value_table(
        document,
        [
            ("Files Reviewed", summary.get("files_reviewed", 0), None),
            ("Lines Added", f"+{summary.get('lines_added', 0)}", COLOR_GREEN),
            ("Lines Removed", f"-{summary.get('lines_removed', 0)}", COLOR_RED),
            ("Overall Risk", title_case(risk), severity_color(risk)),
            ("Final Recommendation", recommendation, recommendation_color(recommendation)),
        ],
    )

    add_heading(document, "Files Changed", level=1)
    add_files_table(document, report.get("files", []))

    issues = report.get("issues", {})
    add_clean_checks(document, issues)
    add_issue_section(document, "Critical Issues", resolve_issue_items(report, issues.get("critical", [])))
    add_issue_section(document, "Breaking Changes", resolve_issue_items(report, issues.get("breaking_changes", [])))
    add_issue_section(document, "Security Concerns", resolve_issue_items(report, issues.get("security", [])))
    add_issue_section(document, "Code Quality Issues", resolve_issue_items(report, issues.get("code_quality", [])))
    add_issue_section(document, "Testing Issues", resolve_issue_items(report, issues.get("testing", [])))
    add_issue_section(document, "Documentation Issues", resolve_issue_items(report, issues.get("documentation", [])))
    add_issue_section(document, "Maintainability Issues", resolve_issue_items(report, issues.get("maintainability", [])))
    add_issue_section(document, "Unprofessional Language", resolve_issue_items(report, issues.get("vulgarity", [])))
    add_issue_section(document, "Performance Issues", resolve_issue_items(report, issues.get("performance", [])))
    add_issue_section(document, "Improvement Suggestions", resolve_issue_items(report, issues.get("improvements", [])))
    add_file_wise_issues(document, report)

    add_heading(document, "Final Recommendation", level=1)
    paragraph = document.add_paragraph()
    add_colored_run(paragraph, str(recommendation), recommendation_color(recommendation), bold=True, size=16)
    reason = document.add_paragraph()
    add_colored_run(reason, "Reason: ", COLOR_MUTED, bold=True)
    add_colored_run(reason, str(report.get("overall_summary", "No summary available.")), COLOR_DARK)

    if report.get("end_summary"):
        final_summary = document.add_paragraph()
        add_colored_run(final_summary, "Summary: ", COLOR_MUTED, bold=True)
        add_colored_run(final_summary, str(report["end_summary"]), COLOR_DARK)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
